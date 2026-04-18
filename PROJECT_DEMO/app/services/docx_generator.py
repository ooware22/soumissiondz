# -*- coding: utf-8 -*-
"""Generateurs DOCX : memoire technique (7 sections) + 3 formulaires algeriens."""
from __future__ import annotations

import io
from datetime import date
from typing import TYPE_CHECKING

from docx import Document as DocxDocument
from docx.shared import Pt

if TYPE_CHECKING:
    from app.models import AppelOffres, Entreprise, Reference


def _h1(doc, text: str) -> None:
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.size = Pt(14)


def _h2(doc, text: str) -> None:
    doc.add_heading(text, level=2)


def _p(doc, text: str) -> None:
    doc.add_paragraph(text)


def generer_memoire_technique(
    entreprise: "Entreprise",
    ao: "AppelOffres",
    references: list["Reference"],
) -> bytes:
    """Genere un memoire technique DOCX en 7 sections a partir des donnees entreprise + AO."""
    doc = DocxDocument()

    doc.add_heading(f"MEMOIRE TECHNIQUE — {ao.objet}", level=0)
    _p(doc, f"AO ref : {ao.reference or '—'}")
    _p(doc, f"Maitre d'ouvrage : {ao.maitre_ouvrage or '—'}")
    _p(doc, f"Date limite : {ao.date_limite.isoformat() if ao.date_limite else '—'}")
    _p(doc, f"Date du memoire : {date.today().isoformat()}")

    # 1 — Presentation de l'entreprise
    _h1(doc, "1. Presentation de l'entreprise")
    _p(doc, f"Denomination sociale : {entreprise.nom}")
    _p(doc, f"Forme juridique : {entreprise.forme_juridique or '—'}")
    _p(doc, f"Gerant : {entreprise.gerant or '—'}")
    _p(doc, f"Siege social : {entreprise.adresse or '—'}, wilaya de {entreprise.wilaya_nom or entreprise.wilaya_code or '—'}")
    _p(doc, f"NIF : {entreprise.nif or '—'} | NIS : {entreprise.nis or '—'} | RC : {entreprise.rc or '—'}")
    _p(doc, f"Effectif : {entreprise.effectif or '—'} personnes. CA moyen annuel : {entreprise.ca_moyen_da or '—'} DA.")

    # 2 — Qualification et activites
    _h1(doc, "2. Qualifications et domaines d'activite")
    _p(doc, f"Categorie de qualification : {entreprise.qualification_cat or '—'}")
    if entreprise.qualification_expiration:
        _p(doc, f"Validite : jusqu'au {entreprise.qualification_expiration.isoformat()}")
    _p(doc, f"Secteur principal : {entreprise.secteur or '—'}")
    _p(doc, f"Activites : {entreprise.activite or '—'}")

    # 3 — Comprehension du projet
    _h1(doc, "3. Comprehension du projet")
    _p(doc, f"L'appel d'offres {ao.reference or ''} porte sur : {ao.objet}.")
    if ao.wilaya_code:
        _p(doc, f"Lieu d'execution : wilaya {ao.wilaya_code}.")
    if ao.budget_estime_da:
        _p(doc, f"Enveloppe budgetaire estimee : {ao.budget_estime_da:,} DA.".replace(",", " "))
    _p(doc, "Notre entreprise a bien identifie les exigences techniques et administratives du cahier des charges et s'engage a les respecter integralement.")

    # 4 — Moyens mis en oeuvre
    _h1(doc, "4. Moyens humains et materiels")
    _p(doc, f"Equipe dediee : encadrement technique et {entreprise.effectif or 'une equipe'} operateurs.")
    _p(doc, "Materiel disponible : engins de chantier, vehicules utilitaires, outillage specialise conforme aux exigences du projet.")
    _p(doc, "La liste detaillee des moyens materiels est jointe en annexe.")

    # 5 — Methodologie d'execution
    _h1(doc, "5. Methodologie d'execution")
    _p(doc, "Phase 1 — Installation de chantier et preparation : 5% du delai.")
    _p(doc, "Phase 2 — Travaux preparatoires (terrassement, ouvrages enterres) : 25% du delai.")
    _p(doc, "Phase 3 — Execution principale : 55% du delai.")
    _p(doc, "Phase 4 — Finitions et reprises : 10% du delai.")
    _p(doc, "Phase 5 — Repli, nettoyage, reception : 5% du delai.")

    # 6 — References
    _h1(doc, "6. References de marches similaires")
    if references:
        for r in references:
            _p(doc, f"• {r.objet} — {r.maitre_ouvrage or '—'} ({r.annee or '—'}) — Montant : {r.montant_da:,} DA.".replace(",", " ") if r.montant_da else
                   f"• {r.objet} — {r.maitre_ouvrage or '—'} ({r.annee or '—'})")
    else:
        _p(doc, "Les references detaillees sont jointes au dossier administratif.")

    # 7 — Engagements
    _h1(doc, "7. Engagements de l'entreprise")
    _p(doc, "L'entreprise s'engage a :")
    _p(doc, "• respecter strictement le delai d'execution figurant dans la lettre de soumission ;")
    _p(doc, "• se conformer aux regles de l'art et aux normes algeriennes en vigueur ;")
    _p(doc, "• appliquer le plan d'hygiene et de securite sur chantier (HSE) ;")
    _p(doc, "• mobiliser une equipe experimentee sous la supervision d'un cadre qualifie.")

    # Signature
    doc.add_paragraph("")
    _p(doc, f"Fait a {entreprise.wilaya_nom or ''}, le {date.today().strftime('%d/%m/%Y')}")
    _p(doc, f"Le gerant : {entreprise.gerant or '—'}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# --------------------------------------------------------------------------
# 3 formulaires officiels algeriens (declaration probite, lettre soumission, decl a souscrire)
# --------------------------------------------------------------------------


def generer_declaration_probite(entreprise: "Entreprise", ao: "AppelOffres") -> bytes:
    doc = DocxDocument()
    doc.add_heading("DECLARATION DE PROBITE", level=0)
    _p(doc, f"Je soussigne {entreprise.gerant or '[Gerant]'}, agissant en qualite de representant legal de la societe {entreprise.nom}, forme juridique {entreprise.forme_juridique or '—'}, NIF {entreprise.nif or '—'}, RC {entreprise.rc or '—'}, ayant son siege a {entreprise.adresse or '—'}, wilaya de {entreprise.wilaya_nom or '—'}.")
    doc.add_paragraph("")
    _p(doc, f"Dans le cadre de la soumission a l'appel d'offres n° {ao.reference or '—'} relatif a : {ao.objet}, lance par {ao.maitre_ouvrage or '—'},")
    doc.add_paragraph("")
    _p(doc, "Declare sur l'honneur :")
    _p(doc, "1. N'avoir pas fait l'objet de condamnation penale definitive pour des faits relevant de la probite professionnelle.")
    _p(doc, "2. N'avoir pas eu recours a des pratiques de corruption, de concussion, de trafic d'influence ou toute autre pratique assimilable.")
    _p(doc, "3. N'avoir pas fait l'objet d'un jugement pour faillite, liquidation judiciaire ou redressement sans plan de continuation.")
    _p(doc, "4. Ne pas etre en etat de conflit d'interet avec le maitre de l'ouvrage ou tout autre membre de la commission d'ouverture.")
    _p(doc, "5. Etre en regle vis-a-vis des obligations fiscales et sociales.")
    doc.add_paragraph("")
    _p(doc, f"Fait a {entreprise.wilaya_nom or ''}, le {date.today().strftime('%d/%m/%Y')}")
    _p(doc, "Signature du representant legal :")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generer_lettre_soumission(entreprise: "Entreprise", ao: "AppelOffres") -> bytes:
    doc = DocxDocument()
    doc.add_heading("LETTRE DE SOUMISSION", level=0)
    _p(doc, f"Je soussigne {entreprise.gerant or '[Gerant]'}, agissant au nom et pour le compte de :")
    _p(doc, f"Denomination : {entreprise.nom}")
    _p(doc, f"Forme juridique : {entreprise.forme_juridique or '—'}")
    _p(doc, f"Siege social : {entreprise.adresse or '—'}, wilaya de {entreprise.wilaya_nom or '—'}")
    _p(doc, f"NIF : {entreprise.nif or '—'} | NIS : {entreprise.nis or '—'} | RC : {entreprise.rc or '—'}")
    _p(doc, f"Telephone : {entreprise.telephone or '—'} | Email : {entreprise.email or '—'}")
    doc.add_paragraph("")
    _p(doc, f"Apres avoir pris connaissance du cahier des charges de l'appel d'offres n° {ao.reference or '—'} portant sur : {ao.objet}, lance par {ao.maitre_ouvrage or '—'},")
    doc.add_paragraph("")
    _p(doc, "Je m'engage, conformement aux clauses et conditions du cahier des charges, a executer les prestations objet dudit appel d'offres moyennant le prix figurant au bordereau des prix unitaires et au detail quantitatif et estimatif joints a la presente offre.")
    doc.add_paragraph("")
    _p(doc, "Le delai d'execution propose est conforme a celui fixe par le cahier des charges.")
    doc.add_paragraph("")
    _p(doc, f"Fait a {entreprise.wilaya_nom or ''}, le {date.today().strftime('%d/%m/%Y')}")
    _p(doc, "Signature et cachet du soumissionnaire :")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generer_declaration_a_souscrire(entreprise: "Entreprise", ao: "AppelOffres") -> bytes:
    doc = DocxDocument()
    doc.add_heading("DECLARATION A SOUSCRIRE", level=0)
    _p(doc, f"Je soussigne {entreprise.gerant or '[Gerant]'}, agissant en qualite de representant legal de {entreprise.nom}, declare :")
    doc.add_paragraph("")
    _p(doc, f"1. Identite : {entreprise.nom}, {entreprise.forme_juridique or '—'}")
    _p(doc, f"2. NIF : {entreprise.nif or '—'} | NIS : {entreprise.nis or '—'} | RC : {entreprise.rc or '—'}")
    _p(doc, f"3. Siege : {entreprise.adresse or '—'}, wilaya de {entreprise.wilaya_nom or '—'}")
    _p(doc, f"4. Coordonnees : {entreprise.telephone or '—'} / {entreprise.email or '—'}")
    _p(doc, f"5. Categorie de qualification : {entreprise.qualification_cat or 'Non qualifie'}")
    _p(doc, f"6. Effectif : {entreprise.effectif or '—'} personnes")
    _p(doc, f"7. CA moyen annuel : {entreprise.ca_moyen_da or '—'} DA")
    doc.add_paragraph("")
    _p(doc, f"Toutes ces informations sont declarees pour valoir dans le cadre de l'appel d'offres n° {ao.reference or '—'} portant sur : {ao.objet}.")
    doc.add_paragraph("")
    _p(doc, "Je certifie sur l'honneur l'exactitude des informations ci-dessus.")
    doc.add_paragraph("")
    _p(doc, f"Fait a {entreprise.wilaya_nom or ''}, le {date.today().strftime('%d/%m/%Y')}")
    _p(doc, "Signature et cachet :")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


FORMULAIRES = {
    "declaration_probite": generer_declaration_probite,
    "lettre_soumission": generer_lettre_soumission,
    "declaration_a_souscrire": generer_declaration_a_souscrire,
}
